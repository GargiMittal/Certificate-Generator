import pandas as pd
from docx import Document
from copy import deepcopy

# Load the CSV file
data = pd.read_csv('farewell.csv')

# Load the template
template = Document('test.docx')

# Create a new document to hold all certificates
certificates_doc = Document()

# Create certificates for each entry in the CSV
for index, row in data.iterrows():
    # Create a new certificate using a deep copy of the template content
    new_certificate = Document()

    # Deep copy the content from the template
    for element in template.element.body:
        new_certificate.element.body.append(deepcopy(element))

    # Find and replace placeholders in the new certificate
    for paragraph in new_certificate.paragraphs:
        if "_______________________" in paragraph.text:
            paragraph.text = paragraph.text.replace("_______________________", row['Name'])
        if "____" in paragraph.text:
            paragraph.text = paragraph.text.replace("____", row['Year'])

    # Add the new certificate to the main document
    for element in new_certificate.element.body:
        certificates_doc.element.body.append(deepcopy(element))

# Save the final document with all certificates
certificates_doc.save('all_certificates.docx')

print("Certificates created successfully!")
